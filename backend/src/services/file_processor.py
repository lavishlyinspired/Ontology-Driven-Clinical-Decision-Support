"""
File Processing Service for Patient Report Analysis

Handles PDF parsing, clinical note extraction, and structured data extraction
from various medical document formats.
"""

import re
import json
from typing import Dict, Any, List, Optional, BinaryIO
from pathlib import Path
import PyPDF2
import docx
from io import BytesIO


class FileProcessor:
    """Process medical documents (PDF, DOCX, TXT) and extract patient data."""
    
    # Clinical data patterns
    PATIENT_PATTERNS = {
        'mrn': r'(?:MRN|Medical Record Number)[:\s]+([A-Z0-9-]+)',
        'name': r'(?:Patient Name|Name)[:\s]+([A-Za-z\s,]+)',
        'dob': r'(?:DOB|Date of Birth)[:\s]+(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})',
        'age': r'(?:Age)[:\s]+(\d{1,3})',
        'sex': r'(?:Sex|Gender)[:\s]+(Male|Female|M|F)',
    }
    
    CLINICAL_PATTERNS = {
        'diagnosis': r'(?:Diagnosis|Dx|Primary Diagnosis)[:\s]+([^\n]+)',
        'stage': r'(?:Stage|TNM Stage)[:\s]+(I{1,3}[ABC]?|IV|Limited|Extensive)',
        'histology': r'(?:Histology|Cell Type)[:\s]+([^\n]+)',
        'performance_status': r'(?:ECOG|Performance Status|PS)[:\s]+(\d)',
        'smoking_history': r'(?:Smoking|Tobacco)[:\s]+([^\n]+)',
    }
    
    BIOMARKER_PATTERNS = {
        'egfr': r'EGFR[:\s]*(Ex19del|Ex20ins|L858R|T790M|positive|negative|[+\-])',
        'alk': r'ALK[:\s]*(positive|negative|fusion detected|rearrangement|[+\-])',
        'ros1': r'ROS1[:\s]*(positive|negative|fusion|[+\-])',
        'pdl1': r'PD-?L1[:\s]*(\d{1,3}%?)',
        'kras': r'KRAS[:\s]*(G12C|G12D|G12V|positive|negative|[+\-])',
        'braf': r'BRAF[:\s]*(V600E|positive|negative|[+\-])',
    }
    
    COMORBIDITY_PATTERNS = {
        'copd': r'(?:COPD|chronic obstructive pulmonary disease)',
        'diabetes': r'(?:diabetes|DM|T2DM)',
        'hypertension': r'(?:hypertension|HTN|high blood pressure)',
        'heart_disease': r'(?:CAD|coronary artery disease|CHF|heart failure)',
        'kidney_disease': r'(?:CKD|chronic kidney disease|renal insufficiency)',
    }
    
    def __init__(self):
        """Initialize file processor."""
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
    def process_file(self, file: BinaryIO, filename: str) -> Dict[str, Any]:
        """
        Process uploaded file and extract patient data.
        
        Args:
            file: File-like object (from FastAPI UploadFile)
            filename: Original filename
            
        Returns:
            Extracted patient data dictionary
        """
        file_ext = Path(filename).suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {file_ext}. Supported: {self.supported_formats}")
        
        # Extract text based on format
        if file_ext == '.pdf':
            text = self._extract_from_pdf(file)
        elif file_ext in ['.docx', '.doc']:
            text = self._extract_from_docx(file)
        else:  # .txt
            text = file.read().decode('utf-8')
        
        # Extract structured data
        patient_data = self._extract_patient_data(text)
        
        return {
            'raw_text': text,
            'extracted_data': patient_data,
            'file_info': {
                'filename': filename,
                'format': file_ext,
                'size_bytes': len(text)
            }
        }
    
    def _extract_from_pdf(self, file: BinaryIO) -> str:
        """Extract text from PDF file."""
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return text
    
    def _extract_from_docx(self, file: BinaryIO) -> str:
        """Extract text from DOCX file."""
        doc = docx.Document(file)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
        
        return text
    
    def _extract_patient_data(self, text: str) -> Dict[str, Any]:
        """Extract structured patient data from text using patterns."""
        extracted = {
            'demographics': {},
            'clinical': {},
            'biomarkers': {},
            'comorbidities': []
        }
        
        # Extract demographics
        for key, pattern in self.PATIENT_PATTERNS.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                extracted['demographics'][key] = match.group(1).strip()
        
        # Extract clinical data
        for key, pattern in self.CLINICAL_PATTERNS.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                extracted['clinical'][key] = match.group(1).strip()
        
        # Extract biomarkers
        for key, pattern in self.BIOMARKER_PATTERNS.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                value = match.group(1).strip()
                # Normalize boolean values
                if value.lower() in ['positive', '+']:
                    value = 'positive'
                elif value.lower() in ['negative', '-']:
                    value = 'negative'
                extracted['biomarkers'][key] = value
        
        # Extract comorbidities
        for key, pattern in self.COMORBIDITY_PATTERNS.items():
            if re.search(pattern, text, re.IGNORECASE):
                extracted['comorbidities'].append(key)
        
        return extracted
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """
        Extract common clinical report sections.
        
        Sections: History, Physical Exam, Imaging, Pathology, Assessment, Plan
        """
        sections = {}
        
        section_patterns = {
            'history': r'(?:HISTORY|CLINICAL HISTORY|HPI)[:\s]+(.*?)(?=\n[A-Z]{4,}|\Z)',
            'physical_exam': r'(?:PHYSICAL EXAM|EXAMINATION)[:\s]+(.*?)(?=\n[A-Z]{4,}|\Z)',
            'imaging': r'(?:IMAGING|RADIOLOGY|CT|PET)[:\s]+(.*?)(?=\n[A-Z]{4,}|\Z)',
            'pathology': r'(?:PATHOLOGY|BIOPSY|HISTOLOGY)[:\s]+(.*?)(?=\n[A-Z]{4,}|\Z)',
            'assessment': r'(?:ASSESSMENT|IMPRESSION)[:\s]+(.*?)(?=\n[A-Z]{4,}|\Z)',
            'plan': r'(?:PLAN|RECOMMENDATION)[:\s]+(.*?)(?=\n[A-Z]{4,}|\Z)',
        }
        
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                sections[section_name] = match.group(1).strip()
        
        return sections
    
    def validate_extracted_data(self, data: Dict[str, Any]) -> Dict[str, List[str]]:
        """
        Validate extracted data and return validation results.
        
        Returns:
            Dictionary with 'errors' and 'warnings' lists
        """
        errors = []
        warnings = []
        
        extracted = data.get('extracted_data', {})
        
        # Required fields
        demographics = extracted.get('demographics', {})
        clinical = extracted.get('clinical', {})
        
        if not demographics.get('age') and not demographics.get('dob'):
            errors.append("Missing age or date of birth")
        
        if not demographics.get('sex'):
            warnings.append("Missing sex/gender information")
        
        if not clinical.get('diagnosis'):
            errors.append("Missing primary diagnosis")
        
        if not clinical.get('stage'):
            warnings.append("Missing disease stage")
        
        # Validate biomarker data
        biomarkers = extracted.get('biomarkers', {})
        if not biomarkers:
            warnings.append("No biomarker data found")
        
        return {
            'errors': errors,
            'warnings': warnings,
            'is_valid': len(errors) == 0
        }
    
    def format_for_lca(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Format extracted data into LCA-compatible patient record.
        
        Returns:
            Patient data in format expected by LCA workflow
        """
        demographics = extracted_data.get('demographics', {})
        clinical = extracted_data.get('clinical', {})
        biomarkers = extracted_data.get('biomarkers', {})
        comorbidities = extracted_data.get('comorbidities', [])
        
        # Map to LCA format
        lca_patient = {
            'patient_id': demographics.get('mrn', 'UNKNOWN'),
            'demographics': {
                'age': int(demographics.get('age', 0)) if demographics.get('age') else None,
                'sex': self._normalize_sex(demographics.get('sex', '')),
            },
            'diagnosis': {
                'cancer_type': self._extract_cancer_type(clinical.get('diagnosis', '')),
                'stage': clinical.get('stage', ''),
                'histology': clinical.get('histology', ''),
            },
            'biomarkers': biomarkers,
            'comorbidities': comorbidities,
            'performance_status': clinical.get('performance_status'),
        }
        
        return lca_patient
    
    def _extract_cancer_type(self, diagnosis: str) -> str:
        """Extract NSCLC/SCLC from diagnosis text."""
        diagnosis_lower = diagnosis.lower()
        
        if 'small cell' in diagnosis_lower or 'sclc' in diagnosis_lower:
            return 'SCLC'
        elif 'non-small cell' in diagnosis_lower or 'nsclc' in diagnosis_lower:
            return 'NSCLC'
        elif any(hist in diagnosis_lower for hist in ['adenocarcinoma', 'squamous', 'large cell']):
            return 'NSCLC'
        
        return 'UNKNOWN'
    
    def _normalize_sex(self, sex: str) -> str:
        """Normalize sex to M/F."""
        sex_upper = sex.upper()
        if sex_upper in ['M', 'MALE']:
            return 'M'
        elif sex_upper in ['F', 'FEMALE']:
            return 'F'
        return ''


class ClinicalNoteParser:
    """Parse unstructured clinical notes using advanced NLP."""
    
    def __init__(self):
        """Initialize clinical note parser."""
        self.file_processor = FileProcessor()
    
    def parse_clinical_note(self, text: str) -> Dict[str, Any]:
        """
        Parse free-text clinical note.
        
        Uses section detection and entity extraction.
        """
        # Extract sections
        sections = self.file_processor.extract_sections(text)
        
        # Extract patient data
        patient_data = self.file_processor._extract_patient_data(text)
        
        # Combine
        return {
            'sections': sections,
            'structured_data': patient_data,
            'metadata': {
                'note_length': len(text),
                'sections_found': list(sections.keys()),
                'entities_extracted': self._count_entities(patient_data)
            }
        }
    
    def _count_entities(self, patient_data: Dict[str, Any]) -> Dict[str, int]:
        """Count extracted entities by category."""
        return {
            'demographics': len(patient_data.get('demographics', {})),
            'clinical': len(patient_data.get('clinical', {})),
            'biomarkers': len(patient_data.get('biomarkers', {})),
            'comorbidities': len(patient_data.get('comorbidities', []))
        }
